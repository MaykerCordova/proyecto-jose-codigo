"""
generar_sustento.py — Documento Word de sustento técnico completo
Tarjetas Comprometidas N7 Débito — Scotiabank Peru

Ejecutar después de 3_ejecutar_ml.bat:
    python scripts/generar_sustento.py

Requiere:
    pip install python-docx

Output:
    output/sustento_scoring_TARJETAS_COMPROMETIDAS_N7.docx
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from config import BASE_DIR, ANALISIS_NOMBRE

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌  Falta python-docx. Instala con:  pip install python-docx")
    sys.exit(1)

OUTPUT_DOC = BASE_DIR / "output" / f"sustento_scoring_{ANALISIS_NOMBRE}.docx"

# ─── Colores corporativos ───────────────────────────────────────────────────
ROJO      = RGBColor(0xC0, 0x00, 0x00)
AZUL_OSC  = RGBColor(0x1F, 0x35, 0x64)
GRIS      = RGBColor(0x59, 0x59, 0x59)
BLANCO    = RGBColor(0xFF, 0xFF, 0xFF)
ROJO_SUAV = RGBColor(0xFF, 0xCC, 0xCC)
VERDE_SUV = RGBColor(0xCC, 0xFF, 0xCC)
AMARI_SUV = RGBColor(0xFF, 0xFF, 0xCC)

# ─── Helpers ────────────────────────────────────────────────────────────────
def set_cell_bg(cell, color_hex):
    """Pone color de fondo a una celda."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex)
    tcPr.append(shd)

def heading(doc, texto, nivel=1, color=None):
    p = doc.add_heading(texto, level=nivel)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p

def parrafo(doc, texto, bold=False, italic=False, size=10, color=None, align=None):
    p   = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p

def tabla_simple(doc, encabezados, filas, col_widths=None, header_bg="1F3564"):
    """Crea tabla con encabezado coloreado y filas alternas."""
    t = doc.add_table(rows=1 + len(filas), cols=len(encabezados))
    t.style = "Table Grid"

    # Encabezado
    hdr = t.rows[0]
    for i, txt in enumerate(encabezados):
        c = hdr.cells[i]
        c.text = txt
        set_cell_bg(c, header_bg)
        for run in c.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = BLANCO
            run.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Filas
    for r_idx, fila in enumerate(filas):
        row = t.rows[r_idx + 1]
        bg  = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(fila):
            c = row.cells[c_idx]
            c.text = str(val)
            set_cell_bg(c, bg)
            c.paragraphs[0].runs[0].font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

def separador(doc):
    doc.add_paragraph("─" * 80)

# ════════════════════════════════════════════════════════════════════════════
# CONSTRUIR DOCUMENTO
# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── PORTADA ─────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SUSTENTO TÉCNICO")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = AZUL_OSC

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Scoring de Fraude — Tarjetas Débito Comprometidas N7")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = ROJO

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Scotiabank Perú | Prevención de Fraude Digital")
r.font.size = Pt(12); r.font.color.rgb = GRIS

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Período analizado: Diciembre 2025 – Mayo 2026")
r.font.size = Pt(11)

doc.add_paragraph()
tabla_simple(doc,
    ["Universo", "Fraudes", "Tasa fraude", "Monto fraude total", "Clientes con fraude"],
    [["170,944 transacciones", "2,288", "1.34%", "S/ 130,043.23", "387 clientes"]],
    col_widths=[4, 2.5, 3, 4, 4],
    header_bg="C00000"
)
doc.add_page_break()

# ── SECCIÓN 1: CONTEXTO ──────────────────────────────────────────────────────
heading(doc, "1. Contexto del Problema", 1, AZUL_OSC)

parrafo(doc,
    "Las tarjetas de débito comprometidas (código N7/INE7) son tarjetas cuyos datos "
    "fueron robados o clonados. El banco las identifica como comprometidas y las segmenta "
    "para monitoreo especial. El desafío analítico es único:", size=10)

doc.add_paragraph()
tabla_simple(doc,
    ["Actor", "Comportamiento", "Etiqueta en la base"],
    [
        ["Titular legítimo", "Compras espaciadas, montos normales, comercios habituales", "N / G (No fraude)"],
        ["Defraudador",      "Muchas compras seguidas, rápido, antes del bloqueo",        "F (Fraude)"],
    ],
    col_widths=[3.5, 9, 4]
)

doc.add_paragraph()
parrafo(doc,
    "Ambos actores usan la MISMA tarjeta. Por eso el label es a nivel de transacción "
    "individual, no a nivel de tarjeta. Esto hace que las técnicas estadísticas sean "
    "imprescindibles: el ojo humano no puede distinguir una transacción legítima de una "
    "fraudulenta sin analizar el contexto completo de comportamiento.", size=10)

doc.add_paragraph()
parrafo(doc, "Hallazgos contraintuitivos confirmados en el análisis:", bold=True, size=10)
tabla_simple(doc,
    ["Intuición común", "Realidad en este dataset", "Explicación"],
    [
        ["El fraude ocurre de madrugada",  "Tasa TARDE (1.47%) > MADRUGADA (1.11%)",  "El defraudador actúa en horario comercial normal"],
        ["El fraude es en montos altos",   "Montos medianos (S/38-58 fraude vs S/22 legítimo)", "Evita llamar la atención con tickets extremos"],
        ["Multi-país = señal de fraude",   "Multi-país tiene MENOR tasa de fraude",   "El fraude N7 es mayoritariamente LOCAL (Perú)"],
        ["Bolivia = riesgo remoto",        "Bolivia = 100% tasa de fraude (9/9)",     "Regla más precisa: declinar automáticamente"],
    ],
    col_widths=[5, 6, 6]
)
doc.add_page_break()

# ── SECCIÓN 2: METODOLOGÍA ──────────────────────────────────────────────────
heading(doc, "2. Metodología — Dos Enfoques Complementarios", 1, AZUL_OSC)

parrafo(doc,
    "Trabajamos con dos enfoques que se complementan. El primero es el que actualmente "
    "usa el equipo de Monitor. El segundo es el que aporta el análisis de datos.", size=10)

doc.add_paragraph()
heading(doc, "2.1 Enfoque Operativo (Monitor — reglas IF-THEN)", 2, ROJO)

parrafo(doc,
    "Se aplica una condición sobre TODA la base histórica y se mide el impacto:", size=10)
parrafo(doc, '   CONDICIÓN: Si MCC = 5411 Y Monto >= S/50 → Alertar', italic=True, size=10)
parrafo(doc, '   RESULTADO: De las 170,944 txn totales...', italic=True, size=10)
parrafo(doc, '     • Se marcaron X,XXX transacciones', italic=True, size=10)
parrafo(doc, '     • De esas, Y son fraude real  →  Precision = Y/X', italic=True, size=10)
parrafo(doc, '     • Y afectamos Z transacciones legítimas  →  Costo de la regla', italic=True, size=10)

doc.add_paragraph()
parrafo(doc,
    "Ventaja: sencillo de implementar en Monitor. "
    "Limitación: cada regla es independiente, no aprende del conjunto de variables.", size=10)

doc.add_paragraph()
heading(doc, "2.2 Enfoque Machine Learning (modelo estadístico)", 2, ROJO)

parrafo(doc,
    "El modelo aprende de los datos históricos para asignar a CADA transacción "
    "una probabilidad de ser fraude P(fraude) entre 0 y 1. "
    "El proceso tiene 4 pasos:", size=10)

doc.add_paragraph()
tabla_simple(doc,
    ["Paso", "Qué hacemos", "Para qué sirve"],
    [
        ["1. Ingeniería de variables", "Creamos 30+ variables de comportamiento (velocidad, z-scores, flags)", "Darle al modelo señales ricas más allá de los datos crudos"],
        ["2. Train/Test split (80/20)", "Separamos 136,598 txn para entrenar y 34,150 para probar", "Simular cómo funciona el modelo con datos nunca vistos"],
        ["3. Entrenar el modelo",       "El modelo aprende qué combinación de variables predice fraude", "Obtener los pesos (coeficientes) de cada variable"],
        ["4. Evaluar y extrapolear",    "Medimos AUC, KS, deciles en el test set y extrapolamos a la base", "Validar que el modelo generaliza antes de producción"],
    ],
    col_widths=[4, 7, 6]
)

doc.add_paragraph()
parrafo(doc, "¿Por qué separamos en Train y Test?", bold=True, size=10)
parrafo(doc,
    "Si entrenamos y evaluamos en la misma data, el modelo simplemente 'memoriza' "
    "los datos y parece perfecto. Al separar, el TEST es data que el modelo NUNCA vio "
    "durante el entrenamiento — simula exactamente lo que pasaría en producción con "
    "transacciones nuevas.", size=10)

doc.add_paragraph()
parrafo(doc, "Extrapolación a la base completa:", bold=True, size=10)
parrafo(doc,
    "Los resultados del Test (34,150 txn, 458 fraudes) se escalan proporcionalmente "
    "a la base total (170,944 txn, 2,288 fraudes). Por ejemplo, si en Test capturamos "
    "346 fraudes de 458 (75.5%), extrapolamos: 75.5% × 2,288 ≈ 1,727 fraudes "
    "en producción.", size=10)
doc.add_page_break()

# ── SECCIÓN 3: REGRESIÓN LOGÍSTICA ──────────────────────────────────────────
heading(doc, "3. El Modelo — Regresión Logística", 1, AZUL_OSC)

heading(doc, "3.1 ¿Qué es la Regresión Logística?", 2, ROJO)
parrafo(doc,
    "La Regresión Logística es un modelo estadístico que produce una ecuación "
    "matemática. Toma las 30 variables de cada transacción y calcula una probabilidad "
    "de fraude entre 0 y 1:", size=10)

doc.add_paragraph()
parrafo(doc,
    "   P(fraude) = 1 / (1 + e^-Z)",
    bold=True, italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
parrafo(doc,
    "   Z = β₀ + β₁×Variable1 + β₂×Variable2 + ... + β₃₀×Variable30",
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
parrafo(doc,
    "Donde cada β (beta) es el peso que el modelo aprendió para cada variable. "
    "Si β > 0, esa variable SUBE la probabilidad de fraude. Si β < 0, la BAJA.", size=10)

doc.add_paragraph()
heading(doc, "3.2 ¿Qué es el Odds Ratio (OR)?", 2, ROJO)
parrafo(doc,
    "El Odds Ratio = e^β es la forma más intuitiva de interpretar los coeficientes:", size=10)
doc.add_paragraph()
tabla_simple(doc,
    ["OR", "Interpretación", "Ejemplo"],
    [
        ["OR = 2.33", "Esta variable MULTIPLICA por 2.33 el riesgo de fraude", "TRX_TARJETA_24H: más transacciones en 24h = 2.3x más riesgo"],
        ["OR = 1.00", "Esta variable NO afecta el riesgo (neutro)",            "Variable sin poder predictivo"],
        ["OR = 0.50", "Esta variable REDUCE a la mitad el riesgo de fraude",   "ES_SEGURO: transacción 3DS = mitad de riesgo"],
    ],
    col_widths=[2.5, 7.5, 7]
)

doc.add_paragraph()
heading(doc, "3.3 Variables del modelo y su traducción a Monitor", 2, ROJO)
parrafo(doc,
    "A continuación, las variables del modelo ordenadas por importancia, "
    "con su interpretación de negocio y cómo implementarlas en Monitor:", size=10)
doc.add_paragraph()

tabla_simple(doc,
    ["Variable", "OR", "Sube/Baja fraude", "Interpretación", "En Monitor"],
    [
        ["TRX_TARJETA_24H",           "2.33", "↑ Sube", "Más txn en 24h = mayor riesgo. Fraude prom: 12.6 / Legítimo: 1.5", "COUNT(txn_tarjeta, 24h) >= 5"],
        ["FLAG_ECOMMERCE",            "2.07", "↑ Sube", "Transacción online = 2x más riesgo que presencial",                "CANAL = 'CNP' (Card Not Present)"],
        ["TRX_CLIENTE_1H",            "1.45", "↑ Sube", "Más txn del cliente en 1h = más riesgo",                          "COUNT(txn_cliente, 1h) >= 3"],
        ["FLAG_COD_TRX_10",           "1.45", "↑ Sube", "Código de transacción 10 asociado a fraude",                      "COD_TRX = '10'"],
        ["CONCENTRACION_5MIN_1H",     "1.36", "↑ Sube", "Ráfagas concentradas en 5min respecto a 1h",                      "Derivar de velocidades Monitor"],
        ["FLAG_MCC_ALTO_RIESGO",      "1.17", "↑ Sube", "MCC en lista de alto riesgo histórico",                           "MCC IN (5411, 4829, 4121, 4722)"],
        ["GAP_MINUTOS",               "0.68", "↓ Baja", "Más tiempo entre compras = menos riesgo. Gap corto = sospechoso", "TIME_SINCE_LAST_TXN <= 10min → alerta"],
        ["FLAG_COD_TRX_92",           "0.48", "↓ Baja", "Código 92 = reversión/especial, protege contra falsos positivos", "Excluir COD_TRX = '92' de reglas"],
        ["FLAG_MONTO_BAJO",           "0.56", "↓ Baja", "Montos muy bajos tienen menos fraude en este segmento",           "MONTO < S/20 → menor prioridad"],
        ["MNT_CLIENTE_24H",           "0.59", "↓ Baja", "Monto acumulado alto paradójicamente baja el riesgo individual",  "Contexto: cliente de alto gasto"],
        ["ES_SEGURO (3DS)",           "0.77", "↓ Baja", "Transacción autenticada con 3DS tiene menos fraude",              "ECI IN ('05','02') → reducir alerta"],
        ["FLAG_MULTI_PAIS_24H",       "0.74", "↓ Baja", "Multi-país baja el riesgo: fraude N7 es LOCAL (Perú)",            "No usar como señal de riesgo en N7"],
    ],
    col_widths=[4.5, 1.5, 2, 5.5, 4.5]
)

doc.add_paragraph()
parrafo(doc,
    "Nota: FLAG_TRX_EN_USD, ES_TOKENIZADA, ES_TARJETA_PRESENTE y ES_MOTO "
    "son todo ceros en este dataset — Monitor no captura estos datos para N7 débito. "
    "Se mantienen en el modelo pero no aportan coeficientes significativos.", size=9, color=GRIS)
doc.add_page_break()

# ── SECCIÓN 4: RESULTADOS DEL MODELO ────────────────────────────────────────
heading(doc, "4. Resultados del Modelo", 1, AZUL_OSC)

heading(doc, "4.1 Métricas de desempeño", 2, ROJO)
parrafo(doc,
    "El modelo fue entrenado en el 80% de la data y evaluado en el 20% restante. "
    "Los resultados son consistentes entre ambos (sin sobreajuste):", size=10)
doc.add_paragraph()

tabla_simple(doc,
    ["Métrica", "Train", "Test", "Benchmark industria", "Resultado"],
    [
        ["AUC-ROC",      "0.7913", "0.7960", "> 0.75 = bueno / > 0.85 = excelente", "✅ Bueno"],
        ["Gini",         "0.5826", "0.5919", "> 0.40 = bueno / > 0.60 = excelente", "✅ Bueno"],
        ["KS Statistic", "0.4442", "0.4691", "> 0.30 = bueno / > 0.50 = excelente", "✅ Bueno"],
        ["Diferencia Train-Test", "—", "0.0047", "< 0.02 = sin sobreajuste", "✅ Sin sobreajuste"],
    ],
    col_widths=[4, 2, 2, 6, 3]
)

doc.add_paragraph()
parrafo(doc,
    "Nota sobre SCORE_RIESGO: En una versión anterior del modelo, incluimos SCORE_RIESGO "
    "(una variable que suma varios flags individuales que también están en el modelo). "
    "Esto generó multicolinealidad severa — el AUC subió artificialmente a 0.978 pero los "
    "coeficientes de las variables individuales quedaron distorsionados (algunos signos "
    "invertidos). En la versión actual, SCORE_RIESGO se excluye del modelo para mantener "
    "la interpretabilidad. SCORE_RIESGO sí se usa como regla de Monitor (Regla 3).", size=9, color=GRIS)

doc.add_paragraph()
heading(doc, "4.2 Matriz de Confusión — ¿Qué pasa en la práctica?", 2, ROJO)
parrafo(doc,
    "La matriz de confusión muestra las 4 situaciones posibles para cada transacción:", size=10)
doc.add_paragraph()

tabla_simple(doc,
    ["", "MODELO: Aprueba", "MODELO: Declina/Alerta"],
    [
        ["REAL: Legítima", "TN — Legítimas bien aprobadas ✅\n(Test: 24,549 | Base: ~122,745)", "FP — Legítimas afectadas por error ⚠️\n(Test: 9,143 | Base: ~45,715)"],
        ["REAL: Fraude",   "FN — Fraudes que escapan ❌\n(Test: 133 | Base: ~665)",             "TP — Fraudes capturados ✅\n(Test: 325 | Base: ~1,625)"],
    ],
    col_widths=[3.5, 7, 7]
)

doc.add_paragraph()
tabla_simple(doc,
    ["Umbral", "Fraudes capturados (TP)", "Fraudes escapan (FN)", "Legítimas afectadas (FP)", "Declina%", "Precision"],
    [
        ["P >= 0.70 — Declinar", "325 test (~1,625 base)", "133 test (~665 base)", "9,143 test (~45,715 base)", "27.7%", "3.4%"],
        ["P >= 0.45 — Revisar",  "325+ test",              "Menos",                "Más",                       "42%",   "2.6%"],
    ],
    col_widths=[4, 3.5, 3.5, 4.5, 2, 2.5]
)

doc.add_paragraph()
heading(doc, "4.3 Análisis por deciles", 2, ROJO)
parrafo(doc,
    "Los deciles ordenan las transacciones de mayor a menor score "
    "y miden cuánto fraude concentra cada grupo del 10%:", size=10)
doc.add_paragraph()

tabla_simple(doc,
    ["Decil", "Score", "Txn", "Fraudes", "Tasa%", "Captura%", "Lift vs base"],
    [
        ["1 (más riesgoso)", "0.65–1.00", "3,415", "209", "6.12%", "45.63%", "4.6x"],
        ["2",               "0.55–0.65", "3,415", "93",  "2.72%", "20.31%", "2.0x"],
        ["3",               "0.49–0.55", "3,415", "24",  "0.70%", "5.24%",  "0.5x"],
        ["4–10",            "0.00–0.49", "23,905","132", "0.55%", "28.82%", "< 1x"],
    ],
    col_widths=[3.5, 2.5, 2, 2, 2, 2.5, 2.5]
)

parrafo(doc,
    "Revisando solo el TOP 20% de transacciones (mayor score) → capturamos el 65.9% del fraude.",
    bold=True, size=10)
doc.add_page_break()

# ── SECCIÓN 5: REGLAS MONITOR ────────────────────────────────────────────────
heading(doc, "5. Las 5 Reglas Monitor — Sustento Completo", 1, AZUL_OSC)

parrafo(doc,
    "Las reglas se calcularon sobre TODA la base (170,944 transacciones), "
    "no solo el test. Esto es equivalente a cómo Monitor evalúa las reglas en producción.", size=10)

doc.add_paragraph()
# Tabla resumen
heading(doc, "5.1 Tabla Resumen de Impacto", 2, ROJO)
tabla_simple(doc,
    ["Regla", "Condición", "Fraudes", "Recall%", "Legítimas", "Monto fraude", "Precision", "Ratio FP/TP"],
    [
        ["R4 — Bolivia",       "País = BO",                         "9",     "0.4%",  "0",     "S/ 679",    "100.0%", "0.0x"],
        ["R5 — Ráfaga 5min",  "3+ txn en 5 min",                   "182",   "8.0%",  "395",   "S/ 1,677",  "31.5%",  "2.2x"],
        ["R1 — Velocidad",    "5+ txn 24h Y GAP <= 10min",         "306",   "13.4%", "3,424", "S/ 4,651",  "8.2%",   "11.2x"],
        ["R2 — MCC riesgo",   "MCC {5411,4829,4121} Y monto>=S/50","546",   "23.9%", "6,415", "S/ 54,262", "7.8%",   "11.7x"],
        ["R3 — Score >= 7",   "SCORE_RIESGO >= 7",                 "1,072", "46.9%", "9,315", "S/ 44,859", "10.3%",  "8.7x"],
    ],
    col_widths=[3.2, 4.5, 1.8, 1.8, 2.2, 2.8, 2.2, 2.2]
)

doc.add_paragraph()
parrafo(doc, "Ratio FP/TP = por cada fraude capturado, cuántas transacciones legítimas afectamos.",
        size=9, color=GRIS, italic=True)

# Cada regla con su sustento
doc.add_paragraph()
heading(doc, "5.2 Sustento por Regla", 2, ROJO)

# REGLA 4
heading(doc, "REGLA 4 — País Bolivia  |  PRIORIDAD: ALTA — Implementar inmediatamente", 3, AZUL_OSC)
tabla_simple(doc,
    ["Métrica", "Valor", "Interpretación"],
    [
        ["Transacciones marcadas", "9",          "Solo 9 txn en todo el período"],
        ["Fraudes capturados",     "9 (100%)",   "Todas las txn de Bolivia son fraude"],
        ["Legítimas afectadas",    "0 (0.00%)",  "Cero clientes buenos afectados"],
        ["Monto fraude protegido", "S/ 678.75",  "Monto total de las 9 txn fraudulentas"],
        ["Monto legítimo afectado","S/ 0.00",    "No hay costo para el banco"],
        ["Clientes fraude",        "1",          "Un solo defraudador identificado"],
        ["Precision",              "100%",       "De cada 10 declinadas, 10 son fraude real"],
    ],
    col_widths=[5, 3, 9]
)
parrafo(doc,
    "Sustento: En el período dic 2025 – may 2026, el 100% de las transacciones "
    "originadas en Bolivia resultaron ser fraude. No existe ningún cliente legítimo "
    "del banco que haya realizado transacciones desde Bolivia en este período. "
    "La regla es perfecta: declinar automáticamente.", size=10)

doc.add_paragraph()
# REGLA 5
heading(doc, "REGLA 5 — Ráfaga en 5 Minutos  |  PRIORIDAD: ALTA — Mejor balance captura/costo", 3, AZUL_OSC)
tabla_simple(doc,
    ["Métrica", "Valor", "Interpretación"],
    [
        ["Transacciones marcadas", "577",        "0.34% del total"],
        ["Fraudes capturados",     "182 (8.0%)", "8 de cada 100 fraudes totales"],
        ["Legítimas afectadas",    "395 (0.23%)","Solo 0.23% de las txn buenas"],
        ["Monto fraude protegido", "S/ 1,676.81",""],
        ["Monto legítimo afectado","S/ 12,359.89",""],
        ["Clientes fraude",        "6",           "6 defraudadores usando ráfagas"],
        ["Clientes legítimos",     "166",         "166 clientes buenos con transacciones rápidas"],
        ["Precision",              "31.5%",       "De cada 10 alertados, 3 son fraude real"],
        ["Recall",                 "8.0%",        "Captura el 8% del fraude total"],
    ],
    col_widths=[5, 3, 9]
)
parrafo(doc,
    "Sustento: Cuando una tarjeta realiza 3+ transacciones en 5 minutos, la tasa de fraude "
    "sube al 31.5% (vs 1.34% de la base). Es la señal más precisa del dataset. "
    "El defraudador intenta hacer el mayor número de compras posible antes de que "
    "el banco bloquee la tarjeta. Por eso compra en ráfaga. "
    "Ratio FP/TP = 2.2: por cada fraude capturado afectamos solo 2.2 transacciones legítimas "
    "(el mejor ratio de todas las reglas).", size=10)

doc.add_paragraph()
# REGLA 3
heading(doc, "REGLA 3 — Score de Riesgo >= 7  |  PRIORIDAD: ALTA — Mayor captura individual", 3, AZUL_OSC)
tabla_simple(doc,
    ["Métrica", "Valor", "Interpretación"],
    [
        ["Transacciones marcadas", "10,387",       "6.1% del total"],
        ["Fraudes capturados",     "1,072 (46.9%)", "Casi la mitad del fraude total"],
        ["Legítimas afectadas",    "9,315 (5.52%)", "5.52% de las txn buenas"],
        ["Monto fraude protegido", "S/ 44,859.03",  "34.5% del monto total de fraude"],
        ["Monto legítimo afectado","S/ 752,672.56", ""],
        ["Clientes fraude",        "218",            "218 defraudadores con score alto"],
        ["Clientes legítimos",     "3,149",          "Clientes con múltiples alertas activas"],
        ["Precision",              "10.3%",          "De cada 10 alertados, 1 es fraude real"],
        ["Recall",                 "46.9%",          "Captura el 46.9% de todo el fraude"],
    ],
    col_widths=[5, 3.5, 8.5]
)
parrafo(doc,
    "Sustento: SCORE_RIESGO es una variable compuesta que suma múltiples alertas "
    "simultáneas (ráfaga de velocidad, MCC de riesgo, multi-país, ecommerce, etc.). "
    "Un score de 7+ significa que VARIAS señales de alerta están activas al mismo tiempo "
    "en una sola transacción. El fraude tiene score promedio de 6.84 vs 3.98 del legítimo. "
    "Esta regla sola captura casi la mitad del fraude.", size=10)

doc.add_paragraph()
# REGLA 2
heading(doc, "REGLA 2 — MCC Alto Riesgo  |  PRIORIDAD: MEDIA — Mayor monto de fraude protegido", 3, AZUL_OSC)
tabla_simple(doc,
    ["Métrica", "Valor", "Interpretación"],
    [
        ["Transacciones marcadas", "6,961",        "4.1% del total"],
        ["Fraudes capturados",     "546 (23.9%)",  "Más de 1 de cada 5 fraudes"],
        ["Legítimas afectadas",    "6,415 (3.80%)","3.80% de las txn buenas"],
        ["Monto fraude protegido", "S/ 54,262.30", "41.7% del monto total de fraude"],
        ["Monto legítimo afectado","S/ 2,084,275.63","Alto costo en monto legítimo"],
        ["Clientes fraude",        "139",           "139 defraudadores en MCCs de riesgo"],
        ["Clientes legítimos",     "1,732",         "Clientes que compran en esos comercios"],
        ["Precision",              "7.8%",          "De cada 10 alertados, 0.78 son fraude"],
        ["Recall",                 "23.9%",         "Captura el 23.9% del fraude total"],
    ],
    col_widths=[5, 4, 8]
)
parrafo(doc,
    "Sustento por MCC: 5411 Supermercados (6.58% fraude, 3,905 txn), "
    "4829 Wire transfers (5.47% fraude, 6,856 txn), "
    "4722 Agencias de viaje (4.26%, 493 txn), "
    "4121 Taxis/Uber (2.50% fraude, 30,409 txn — mayor volumen). "
    "El filtro de monto >= S/50 reduce falsos positivos de tickets muy bajos.", size=10)

doc.add_paragraph()
# REGLA 1
heading(doc, "REGLA 1 — Velocidad Extrema  |  PRIORIDAD: MEDIA — Complementa R5", 3, AZUL_OSC)
tabla_simple(doc,
    ["Métrica", "Valor", "Interpretación"],
    [
        ["Transacciones marcadas", "3,730",        "2.18% del total"],
        ["Fraudes capturados",     "306 (13.4%)",  ""],
        ["Legítimas afectadas",    "3,424 (2.03%)",""],
        ["Monto fraude protegido", "S/ 4,651.40",  "Montos bajos por txn individual"],
        ["Monto legítimo afectado","S/ 191,038.18",""],
        ["Clientes fraude",        "31",            "31 defraudadores con velocidad extrema"],
        ["Clientes legítimos",     "449",           "Clientes con uso intensivo de la tarjeta"],
        ["Precision",              "8.2%",          "De cada 10 alertados, 0.82 son fraude"],
        ["Recall",                 "13.4%",         "Captura el 13.4% del fraude total"],
    ],
    col_widths=[5, 3.5, 8.5]
)
parrafo(doc,
    "Sustento: El fraude promedio realiza 12.6 txn en 24 horas (vs 1.5 del legítimo) "
    "con un gap promedio de 2,087 minutos entre txn (vs 6,718 del legítimo). "
    "Cuando TRX_24H >= 5 Y GAP <= 10 minutos, se detecta al defraudador en plena ráfaga. "
    "Ratio FP/TP = 11.2 (mayor costo que R5). Recomendamos evaluar después de implementar R5.", size=10)
doc.add_page_break()

# ── SECCIÓN 6: CASCADA ──────────────────────────────────────────────────────
heading(doc, "6. Análisis en Cascada — Implementación por Fases", 1, AZUL_OSC)

parrafo(doc,
    "La cascada muestra la captura acumulada si aplicamos las reglas en secuencia. "
    "Las reglas se aplican sobre transacciones NO capturadas por las anteriores "
    "(sin doble conteo):", size=10)
doc.add_paragraph()

tabla_simple(doc,
    ["Fase", "Regla", "Fraudes nuevos", "Acumulado", "% Total", "Legítimas acum.", "Precision acum."],
    [
        ["FASE 1", "R4 — Bolivia",      "9",   "9",     "0.4%",  "0",      "100.0%"],
        ["FASE 1", "R5 — Ráfaga 5min",  "182", "191",   "8.4%",  "395",    "~48%"],
        ["FASE 2", "R3 — Score >= 7",   "781", "1,090", "47.6%", "11,644", "8.6%"],
        ["FASE 3", "R2 — MCC riesgo",   "386", "1,476", "64.5%", "17,699", "7.7%"],
        ["FASE 3", "R1 — Velocidad",    "306", "1,482", "64.8%", "17,699", "7.7%"],
    ],
    col_widths=[2, 3.5, 3, 2.5, 2.5, 3.5, 3]
)

doc.add_paragraph()
parrafo(doc,
    "RECOMENDACIÓN DE IMPLEMENTACIÓN:", bold=True, size=11)
parrafo(doc,
    "FASE 1 — Esta semana: R4 + R5 → Captura 8.4% del fraude, afecta solo 395 legítimas. "
    "Riesgo operativo casi nulo.", size=10)
parrafo(doc,
    "FASE 2 — Siguiente mes: + R3 → Salta a 47.6% del fraude. "
    "Evaluar capacidad del equipo de revisión (9,315 nuevas alertas).", size=10)
parrafo(doc,
    "FASE 3 — Después de validar F2: + R2 → Llega a 64.5% del fraude. "
    "Ajustar umbrales de monto si la afectación al legítimo es alta.", size=10)
doc.add_page_break()

# ── SECCIÓN 7: OTRAS TÉCNICAS ML ────────────────────────────────────────────
heading(doc, "7. Hacia el Machine Learning Automático en Fraude", 1, AZUL_OSC)

parrafo(doc,
    "Lo que construimos es solo el primer paso. La propuesta de valor del análisis "
    "de datos va mucho más allá de las 5 reglas actuales:", size=10)

doc.add_paragraph()
heading(doc, "7.1 Técnicas supervisadas disponibles", 2, ROJO)
tabla_simple(doc,
    ["Técnica", "Cómo funciona", "Ventaja sobre reglas manuales", "Cuándo usarla"],
    [
        ["Regresión Logística\n(actual)", "Ecuación matemática con pesos por variable", "Combina 30 variables simultáneamente, aprende de la data", "Cuando se necesita interpretabilidad (explicar a reguladores)"],
        ["Árboles de Decisión",           "Genera reglas IF-THEN automáticamente desde la data", "Las reglas resultantes son intuitivas y directas", "Para generar reglas Monitor automáticamente"],
        ["Random Forest",                 "Ensambla 100+ árboles y promedia sus predicciones", "Muy robusto, maneja bien el desbalance", "Cuando la precisión importa más que la interpretabilidad"],
        ["XGBoost / Gradient Boosting",   "Árboles secuenciales que corrigen errores del anterior", "Mejor AUC en la mayoría de competencias de fraude", "Cuando se necesita máxima capacidad predictiva"],
        ["Redes Neuronales",              "Capas de neuronas que aprenden patrones complejos", "Detecta patrones no lineales muy sofisticados", "Con grandes volúmenes y features de texto/imagen"],
    ],
    col_widths=[3.5, 4, 5, 5]
)

doc.add_paragraph()
heading(doc, "7.2 Técnicas no supervisadas", 2, ROJO)
tabla_simple(doc,
    ["Técnica", "Para qué sirve en fraude"],
    [
        ["Clustering (K-means, DBSCAN)", "Agrupar transacciones por comportamiento similar — detectar grupos sospechosos sin etiquetas previas"],
        ["Isolation Forest",            "Detectar anomalías: transacciones que 'no encajan' con el patrón normal del cliente"],
        ["Autoencoders",                "Aprender el comportamiento normal y flagear lo que se desvía significativamente"],
        ["Graph Analytics",             "Detectar redes de fraude: defraudadores que comparten comercios, IPs o patrones de tiempo"],
    ],
    col_widths=[4.5, 13]
)

doc.add_paragraph()
heading(doc, "7.3 Metodología estándar — Del problema a la solución", 2, ROJO)
tabla_simple(doc,
    ["Paso", "En este proyecto", "En general"],
    [
        ["1. Definir el problema",        "¿Qué txn son fraude en tarjetas comprometidas N7?",               "¿Qué quiero predecir/detectar?"],
        ["2. Recopilar y limpiar datos",  "170,944 txn, 268 columnas, dic 2025 – may 2026",                 "Data histórica etiquetada"],
        ["3. Ingeniería de variables",    "30 variables de comportamiento (velocidad, z-scores, flags)",     "Crear señales ricas desde datos crudos"],
        ["4. Seleccionar técnica",        "Regresión Logística (AUC >= 0.75 → no se necesita XGBoost)",     "Según interpretabilidad vs performance"],
        ["5. Entrenar y validar",         "Train 80% / Test 20% estratificado — AUC 0.796, KS 0.469",      "Métricas en datos nunca vistos"],
        ["6. Interpretar resultados",     "Odds Ratios, deciles, umbrales operativos",                       "¿Qué dice el modelo? ¿Es confiable?"],
        ["7. Traducir a acciones",        "5 reglas Monitor con sustento de impacto",                        "Reglas, alertas, modelos en producción"],
        ["8. Monitorear en el tiempo",    "PSI por feature (estabilidad), seguimiento de métricas",          "¿El modelo sigue siendo válido?"],
    ],
    col_widths=[3.5, 7, 7]
)

doc.add_paragraph()
parrafo(doc,
    "IMPACTO: Pasar de reglas manuales a machine learning permite:",
    bold=True, size=11)
parrafo(doc, "  • Analizar 30+ variables simultáneamente (imposible manualmente)", size=10)
parrafo(doc, "  • Cuantificar exactamente el costo/beneficio de cada regla", size=10)
parrafo(doc, "  • Detectar patrones contraintuitivos (fraude de día, no de noche)", size=10)
parrafo(doc, "  • Actualizar el modelo periódicamente con nueva data", size=10)
parrafo(doc, "  • Escalar a otros segmentos (ecommerce, tarjetas de crédito, etc.)", size=10)
doc.add_page_break()

# ── SECCIÓN 8: CONCLUSIONES ──────────────────────────────────────────────────
heading(doc, "8. Conclusiones y Próximos Pasos", 1, AZUL_OSC)

parrafo(doc, "Lo que logramos:", bold=True, size=11)
parrafo(doc, "  ✅ Modelo de scoring con AUC=0.796, KS=0.469, sin sobreajuste", size=10)
parrafo(doc, "  ✅ 5 reglas Monitor con sustento cuantitativo completo", size=10)
parrafo(doc, "  ✅ Identificación de las variables más predictivas del fraude N7", size=10)
parrafo(doc, "  ✅ Captura potencial del 64.8% del fraude con las 5 reglas combinadas", size=10)
parrafo(doc, "  ✅ Bolivia: 1 regla perfecta (100% precisión, 0 afectados) lista para implementar", size=10)

doc.add_paragraph()
parrafo(doc, "Próximos pasos sugeridos:", bold=True, size=11)
tabla_simple(doc,
    ["Prioridad", "Acción", "Impacto esperado"],
    [
        ["INMEDIATO",  "Implementar R4 (Bolivia) en Monitor",              "9 fraudes bloqueados, 0 costo"],
        ["INMEDIATO",  "Implementar R5 (Ráfaga 5min) en Monitor",          "182 fraudes, Precision 31.5%"],
        ["1 MES",      "Implementar R3 (Score >= 7) con revisión manual",  "Captura acumulada 47.6%"],
        ["2 MESES",    "Evaluar R2 (MCC) y ajustar umbral de monto",       "Captura acumulada 64.5%"],
        ["3 MESES",    "Reentrenar el modelo con 6 meses adicionales",     "Incorporar nuevos patrones"],
        ["FUTURO",     "Explorar árboles de decisión para generar reglas automáticamente", "Escalabilidad"],
    ],
    col_widths=[3, 7, 7]
)

doc.add_paragraph()
separador(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Scotiabank Perú | Prevención de Fraude Digital | Análisis generado con Python + Machine Learning")
r.font.size = Pt(8)
r.font.color.rgb = GRIS

# ── GUARDAR ──────────────────────────────────────────────────────────────────
OUTPUT_DOC.parent.mkdir(exist_ok=True)
try:
    doc.save(OUTPUT_DOC)
    print(f"✅  Documento guardado: {OUTPUT_DOC}")
except Exception as e:
    print(f"❌  Error al guardar: {e}")
